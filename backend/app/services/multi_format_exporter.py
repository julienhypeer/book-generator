"""
Service multi-formats pour export EPUB, DOCX.
Implémentation simple et directe selon les tests TDD.
"""

import asyncio
import time
import zipfile
from io import BytesIO
from typing import Dict, List, Tuple, Any
from datetime import datetime
import logging
import html
import re
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)

# Imports conditionnels pour les dépendances
try:
    from ebooklib import epub

    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False
    logger.warning("ebooklib not available for EPUB export")

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available for DOCX export")

# App imports after conditional imports to avoid E402
from app.models.project import Project  # noqa: E402
from app.models.chapter import Chapter  # noqa: E402
from app.validators.export import ExportRequest, TemplateType  # noqa: E402


class ExportError(Exception):
    """Exception pour les erreurs d'export."""

    pass


class BaseGenerator(ABC):
    """Classe de base pour tous les générateurs."""

    @abstractmethod
    def generate(
        self, project: Project, chapters: List[Chapter], **kwargs
    ) -> bytes:
        """Génère le fichier d'export."""
        pass

    def sanitize_html(self, html_content: str) -> str:
        """Nettoie et corrige le HTML."""
        if not html_content:
            return ""

        # Corrections basiques HTML
        html_content = html_content.replace("<br>", "<br/>")
        html_content = html_content.replace("<hr>", "<hr/>")

        # Échapper les caractères spéciaux XML
        html_content = html.escape(html_content, quote=False)
        html_content = (
            html_content.replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&amp;", "&")
        )

        return html_content

    def markdown_to_html(self, markdown_content: str) -> str:
        """Conversion Markdown basique vers HTML."""
        if not markdown_content:
            return ""

        # Conversion simple des éléments Markdown les plus courants
        html_content = markdown_content

        # Titres
        html_content = re.sub(
            r"^# (.*)", r"<h1>\1</h1>", html_content, flags=re.MULTILINE
        )
        html_content = re.sub(
            r"^## (.*)", r"<h2>\1</h2>", html_content, flags=re.MULTILINE
        )
        html_content = re.sub(
            r"^### (.*)", r"<h3>\1</h3>", html_content, flags=re.MULTILINE
        )

        # Gras et italique
        html_content = re.sub(
            r"\*\*(.*?)\*\*", r"<strong>\1</strong>", html_content
        )
        html_content = re.sub(r"\*(.*?)\*", r"<em>\1</em>", html_content)

        # Code inline
        html_content = re.sub(r"`(.*?)`", r"<code>\1</code>", html_content)

        # Listes
        html_content = re.sub(
            r"^- (.*)", r"<li>\1</li>", html_content, flags=re.MULTILINE
        )
        html_content = re.sub(
            r"(<li>.*?</li>\n?)+",
            r"<ul>\g<0></ul>",
            html_content,
            flags=re.DOTALL,
        )

        # Citations
        html_content = re.sub(
            r"^> (.*)",
            r"<blockquote><p>\1</p></blockquote>",
            html_content,
            flags=re.MULTILINE,
        )

        # Paragraphes (lignes non vides qui ne sont pas déjà des éléments HTML)
        lines = html_content.split("\n")
        processed_lines = []
        for line in lines:
            line = line.strip()
            if line and not re.match(r"<[^>]+>", line):
                processed_lines.append(f"<p>{line}</p>")
            else:
                processed_lines.append(line)

        return "\n".join(processed_lines)


class EPUBGenerator(BaseGenerator):
    """Générateur EPUB utilisant ebooklib."""

    def generate_epub(
        self, project: Project, chapters: List[Chapter]
    ) -> bytes:
        """Génère un fichier EPUB."""
        if not EPUB_AVAILABLE:
            return self._generate_manual_epub(project, chapters)

        # Créer le livre EPUB
        book = epub.EpubBook()

        # Métadonnées
        book.set_identifier(f"book-{project.id}-{int(time.time())}")
        book.set_title(project.title)
        book.set_language("fr")

        if project.settings:
            if "author" in project.settings:
                book.add_author(project.settings["author"])
            if "description" in project.settings:
                book.add_metadata(
                    "DC", "description", project.settings["description"]
                )

        # Ajouter les chapitres
        chapter_items = []
        for i, chapter in enumerate(chapters, 1):
            chapter_content = self.sanitize_html(chapter.content)
            if "<h1>" not in chapter_content and "<h2>" not in chapter_content:
                chapter_content = (
                    f"<h1>{chapter.title}</h1>\n{chapter_content}"
                )

            epub_chapter = epub.EpubHtml(
                title=chapter.title, file_name=f"chapter{i}.xhtml", lang="fr"
            )
            epub_chapter.content = f"""<?xml version='1.0' encoding='utf-8'?>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{chapter.title}</title>
</head>
<body>
    {chapter_content}
</body>
</html>"""

            book.add_item(epub_chapter)
            chapter_items.append(epub_chapter)

        # Table des matières
        book.toc = chapter_items
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Définir l'ordre de lecture
        book.spine = ["nav"] + chapter_items

        # Générer le fichier EPUB
        output = BytesIO()
        epub.write_epub(output, book)
        return output.getvalue()

    def _generate_manual_epub(
        self, project: Project, chapters: List[Chapter]
    ) -> bytes:
        """Génération EPUB manuelle si ebooklib n'est pas disponible."""
        output = BytesIO()

        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as epub_zip:
            # 1. mimetype (non compressé)
            epub_zip.writestr(
                "mimetype",
                "application/epub+zip",
                compress_type=zipfile.ZIP_STORED,
            )

            # 2. META-INF/container.xml
            container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" \
media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>"""
            epub_zip.writestr("META-INF/container.xml", container_xml)

            # 3. OEBPS/content.opf
            opf_content = self._generate_opf_content(project, chapters)
            epub_zip.writestr("OEBPS/content.opf", opf_content)

            # 4. OEBPS/toc.ncx
            ncx_content = self._generate_ncx_content(project, chapters)
            epub_zip.writestr("OEBPS/toc.ncx", ncx_content)

            # 5. Chapitres XHTML
            for i, chapter in enumerate(chapters, 1):
                chapter_html = self._generate_chapter_xhtml(chapter)
                epub_zip.writestr(f"OEBPS/chapter{i}.xhtml", chapter_html)

        return output.getvalue()

    def _generate_opf_content(
        self, project: Project, chapters: List[Chapter]
    ) -> str:
        """Génère le fichier content.opf."""
        author = (
            project.settings.get("author", "Unknown")
            if project.settings
            else "Unknown"
        )
        description = (
            project.settings.get("description", "") if project.settings else ""
        )

        # Manifest des chapitres
        manifest_items = []
        spine_items = []

        for i, chapter in enumerate(chapters, 1):
            manifest_items.append(
                f'<item id="chapter{i}" href="chapter{i}.xhtml" '
                f'media-type="application/xhtml+xml"/>'
            )
            spine_items.append(f'<itemref idref="chapter{i}"/>')

        opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="bookid" version="2.0">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:title>{html.escape(project.title)}</dc:title>
        <dc:creator>{html.escape(author)}</dc:creator>
        <dc:language>fr</dc:language>
        <dc:identifier id="bookid" opf:scheme="UUID">\
book-{project.id}-{int(time.time())}</dc:identifier>
        <dc:date>{datetime.now().strftime('%Y-%m-%d')}</dc:date>
        <dc:description>{html.escape(description)}</dc:description>
    </metadata>
    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        {chr(10).join(manifest_items)}
    </manifest>
    <spine toc="ncx">
        {chr(10).join(spine_items)}
    </spine>
</package>"""

        return opf

    def _generate_ncx_content(
        self, project: Project, chapters: List[Chapter]
    ) -> str:
        """Génère le fichier toc.ncx."""
        nav_points = []

        for i, chapter in enumerate(chapters, 1):
            nav_points.append(
                f"""
        <navPoint id="chapter{i}" playOrder="{i}">
            <navLabel>
                <text>{html.escape(chapter.title)}</text>
            </navLabel>
            <content src="chapter{i}.xhtml"/>
        </navPoint>"""
            )

        ncx = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx version="2005-1" xml:lang="fr" xmlns="http://www.daisy.org/z3986/2005/ncx/">
    <head>
        <meta name="dtb:uid" content="book-{project.id}"/>
        <meta name="dtb:depth" content="1"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    <docTitle>
        <text>{html.escape(project.title)}</text>
    </docTitle>
    <navMap>
        {chr(10).join(nav_points)}
    </navMap>
</ncx>"""

        return ncx

    def _generate_chapter_xhtml(self, chapter: Chapter) -> str:
        """Génère le XHTML d'un chapitre."""
        content = self.sanitize_html(chapter.content)

        return f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{html.escape(chapter.title)}</title>
    <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
</head>
<body>
    <h1>{html.escape(chapter.title)}</h1>
    {content}
</body>
</html>"""

    def generate(
        self, project: Project, chapters: List[Chapter], **kwargs
    ) -> bytes:
        """Interface BaseGenerator."""
        return self.generate_epub(project, chapters)


class DOCXGenerator(BaseGenerator):
    """Générateur DOCX utilisant python-docx."""

    def generate_docx(
        self, project: Project, chapters: List[Chapter]
    ) -> bytes:
        """Génère un fichier DOCX."""
        if not DOCX_AVAILABLE:
            return self._generate_manual_docx(project, chapters)

        # Créer le document
        doc = Document()

        # Métadonnées
        doc.core_properties.title = project.title
        if project.settings:
            doc.core_properties.author = project.settings.get("author", "")
            doc.core_properties.subject = project.settings.get("subject", "")
            doc.core_properties.comments = project.settings.get(
                "description", ""
            )

        # Ajouter titre principal
        title = doc.add_heading(project.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ajouter les chapitres
        for chapter in chapters:
            # Titre du chapitre
            doc.add_heading(chapter.title, level=1)

            # Contenu du chapitre
            content = self._convert_html_to_docx_content(chapter.content)
            self._add_content_to_doc(doc, content)

            # Saut de page après chaque chapitre (sauf le dernier)
            if chapter != chapters[-1]:
                doc.add_page_break()

        # Sauvegarder en mémoire
        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _convert_html_to_docx_content(
        self, html_content: str
    ) -> List[Dict[str, str]]:
        """Convertit HTML en éléments DOCX structurés."""
        if not html_content:
            return []

        elements = []

        # Parsing HTML simple
        html_content = self.sanitize_html(html_content)

        # Diviser en paragraphes basiques
        lines = html_content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Titres
            if line.startswith("<h1>") and line.endswith("</h1>"):
                text = line.replace("<h1>", "").replace("</h1>", "")
                elements.append({"type": "heading", "level": 1, "text": text})
            elif line.startswith("<h2>") and line.endswith("</h2>"):
                text = line.replace("<h2>", "").replace("</h2>", "")
                elements.append({"type": "heading", "level": 2, "text": text})
            elif line.startswith("<h3>") and line.endswith("</h3>"):
                text = line.replace("<h3>", "").replace("</h3>", "")
                elements.append({"type": "heading", "level": 3, "text": text})
            # Paragraphes
            elif line.startswith("<p>") and line.endswith("</p>"):
                text = line.replace("<p>", "").replace("</p>", "")
                text = self._clean_inline_html(text)
                elements.append({"type": "paragraph", "text": text})
            # Code
            elif "<code>" in line or "<pre>" in line:
                text = re.sub(r"<[^>]+>", "", line)
                elements.append({"type": "code", "text": text})
            # Listes
            elif line.startswith("<li>") and line.endswith("</li>"):
                text = line.replace("<li>", "").replace("</li>", "")
                elements.append({"type": "list_item", "text": text})
            # Citations
            elif "<blockquote>" in line:
                text = re.sub(r"<[^>]+>", "", line)
                elements.append({"type": "quote", "text": text})
            # Paragraphe par défaut
            else:
                text = self._clean_inline_html(line)
                if text:
                    elements.append({"type": "paragraph", "text": text})

        return elements

    def _clean_inline_html(self, text: str) -> str:
        """Nettoie les balises HTML inline."""
        # Conserver le formatage de base
        text = text.replace("<strong>", "**").replace("</strong>", "**")
        text = text.replace("<em>", "*").replace("</em>", "*")
        text = text.replace("<code>", "`").replace("</code>", "`")

        # Supprimer les autres balises
        text = re.sub(r"<[^>]+>", "", text)

        return text

    def _add_content_to_doc(self, doc, content_elements: List[Dict[str, str]]):
        """Ajoute les éléments de contenu au document DOCX."""
        for element in content_elements:
            if element["type"] == "heading":
                doc.add_heading(element["text"], level=element["level"])
            elif element["type"] == "paragraph":
                doc.add_paragraph(element["text"])
            elif element["type"] == "code":
                p = doc.add_paragraph()
                p.style = "Normal"
                run = p.add_run(element["text"])
                run.font.name = "Courier New"
            elif element["type"] == "list_item":
                doc.add_paragraph(f"• {element['text']}")
            elif element["type"] == "quote":
                p = doc.add_paragraph(element["text"])
                p.style = "Quote"

    def _generate_manual_docx(
        self, project: Project, chapters: List[Chapter]
    ) -> bytes:
        """Génération DOCX manuelle si python-docx n'est pas disponible."""
        output = BytesIO()

        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as docx_zip:
            # 1. [Content_Types].xml
            content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
    <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
</Types>"""
            docx_zip.writestr("[Content_Types].xml", content_types)

            # 2. _rels/.rels
            main_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
    <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
</Relationships>"""
            docx_zip.writestr("_rels/.rels", main_rels)

            # 3. word/_rels/document.xml.rels
            doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>"""
            docx_zip.writestr("word/_rels/document.xml.rels", doc_rels)

            # 4. docProps/core.xml
            core_props = self._generate_core_properties(project)
            docx_zip.writestr("docProps/core.xml", core_props)

            # 5. word/document.xml
            document_xml = self._generate_document_xml(project, chapters)
            docx_zip.writestr("word/document.xml", document_xml)

        return output.getvalue()

    def _generate_core_properties(self, project: Project) -> str:
        """Génère docProps/core.xml."""
        author = project.settings.get("author", "") if project.settings else ""
        subject = (
            project.settings.get("subject", "") if project.settings else ""
        )

        return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <dc:title>{html.escape(project.title)}</dc:title>
    <dc:creator>{html.escape(author)}</dc:creator>
    <dc:subject>{html.escape(subject)}</dc:subject>
    <dcterms:created xsi:type="dcterms:W3CDTF">{datetime.now().isoformat()}Z</dcterms:created>
    <dcterms:modified xsi:type="dcterms:W3CDTF">{datetime.now().isoformat()}Z</dcterms:modified>
</cp:coreProperties>"""

    def _generate_document_xml(
        self, project: Project, chapters: List[Chapter]
    ) -> str:
        """Génère word/document.xml."""
        paragraphs = []

        # Titre principal
        paragraphs.append(
            f"""
        <w:p>
            <w:pPr>
                <w:jc w:val="center"/>
            </w:pPr>
            <w:r>
                <w:rPr>
                    <w:b/>
                    <w:sz w:val="32"/>
                </w:rPr>
                <w:t>{html.escape(project.title)}</w:t>
            </w:r>
        </w:p>"""
        )

        # Chapitres
        for chapter in chapters:
            # Titre de chapitre
            paragraphs.append(
                f"""
        <w:p>
            <w:pPr>
                <w:pStyle w:val="Heading1"/>
            </w:pPr>
            <w:r>
                <w:t>{html.escape(chapter.title)}</w:t>
            </w:r>
        </w:p>"""
            )

            # Contenu de chapitre (simplifié)
            content_text = re.sub(r"<[^>]+>", "", chapter.content or "")
            content_text = html.escape(content_text)

            paragraphs.append(
                f"""
        <w:p>
            <w:r>
                <w:t>{content_text}</w:t>
            </w:r>
        </w:p>"""
            )

        return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:body>
        {chr(10).join(paragraphs)}
    </w:body>
</w:document>"""

    def generate(
        self, project: Project, chapters: List[Chapter], **kwargs
    ) -> bytes:
        """Interface BaseGenerator."""
        return self.generate_docx(project, chapters)


class MultiFormatExporter:
    """Service principal pour l'export multi-formats."""

    def __init__(self):
        self.generators = {"epub": EPUBGenerator(), "docx": DOCXGenerator()}

    async def export(
        self,
        project: Project,
        chapters: List[Chapter],
        export_request: ExportRequest,
    ) -> Tuple[bytes, Dict[str, Any]]:
        """Exporte le projet dans le format demandé."""
        start_time = time.time()

        # Validation
        if not chapters:
            raise ExportError("No chapters found in project")

        format_name = export_request.format.lower()
        if format_name not in self.generators:
            raise ExportError(f"Unsupported format: {format_name}")

        # Générer le fichier
        generator = self._get_generator(format_name)
        result_bytes = await self._run_generator(
            generator, project, chapters, export_request
        )

        # Métadonnées
        generation_time = time.time() - start_time
        metadata = {
            "format": format_name,
            "file_size": len(result_bytes),
            "generation_time": generation_time,
            "chapter_count": len(chapters),
            "export_settings": {
                "template": export_request.template,
                "include_toc": export_request.include_toc,
                "include_page_numbers": export_request.include_page_numbers,
                "quality_validation": export_request.quality_validation,
            },
        }

        logger.info(
            f"Export completed: {format_name.upper()}, "
            f"{len(chapters)} chapters, "
            f"{len(result_bytes)} bytes, "
            f"{generation_time:.2f}s"
        )

        return result_bytes, metadata

    async def _run_generator(
        self,
        generator: BaseGenerator,
        project: Project,
        chapters: List[Chapter],
        export_request: ExportRequest,
    ) -> bytes:
        """Exécute le générateur de façon asynchrone."""
        # Exécuter en arrière-plan pour ne pas bloquer
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, generator.generate, project, chapters
        )

    def _get_generator(self, format_name: str) -> BaseGenerator:
        """Récupère le générateur pour le format."""
        if format_name not in self.generators:
            raise ExportError(f"Unsupported format: {format_name}")
        return self.generators[format_name]

    def get_available_templates(self) -> List[TemplateType]:
        """Retourne les templates disponibles."""
        return [
            TemplateType.ROMAN,
            TemplateType.TECHNICAL,
            TemplateType.ACADEMIC,
        ]

    def _get_template_styles(self, template: TemplateType) -> Dict[str, str]:
        """Retourne les styles pour un template."""
        if template == TemplateType.ROMAN:
            return {
                "font_family": "Crimson Text, Georgia, serif",
                "font_size": "12pt",
                "line_height": "1.6",
                "text_align": "justify",
            }
        elif template == TemplateType.TECHNICAL:
            return {
                "font_family": "Source Sans Pro, Arial, sans-serif",
                "font_size": "11pt",
                "line_height": "1.4",
                "text_align": "left",
            }
        elif template == TemplateType.ACADEMIC:
            return {
                "font_family": "Times New Roman, serif",
                "font_size": "12pt",
                "line_height": "2.0",
                "text_align": "justify",
            }
        else:
            return {
                "font_family": "Arial, sans-serif",
                "font_size": "12pt",
                "line_height": "1.5",
                "text_align": "left",
            }
